"""
Document parsing service for PDF, DOCX, and TXT files
"""
import base64
import io
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class ParserService:
    """Service to extract text from various document formats"""
    
    @staticmethod
    def parse_pdf(content: bytes) -> str:
        """Extract text from PDF using pdfplumber (better accuracy)"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                
                reader = PdfReader(io.BytesIO(content))
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
            except Exception as e2:
                logger.error(f"Both PDF parsers failed: {e2}")
                raise ValueError(f"Failed to parse PDF: {str(e2)}")
    
    @staticmethod
    def parse_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_txt(content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try common encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # Last resort: decode with errors ignored
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"TXT parsing failed: {e}")
            raise ValueError(f"Failed to parse TXT: {str(e)}")
    
    @classmethod
    def parse(cls, content: str, file_type: str) -> str:
        """
        Parse document content based on file type
        
        Args:
            content: Base64 encoded file content or plain text
            file_type: File extension (pdf, docx, txt)
        
        Returns:
            Extracted text content
        """
        file_type = file_type.lower()
        
        # Handle plain text input (not base64)
        if file_type == "txt":
            try:
                decoded = base64.b64decode(content)
                return cls.parse_txt(decoded)
            except Exception:
                # Might already be plain text
                return content
        
        # Decode base64 for binary files
        try:
            decoded = base64.b64decode(content)
        except Exception as e:
            raise ValueError(f"Invalid base64 content: {str(e)}")
        
        if file_type == "pdf":
            return cls.parse_pdf(decoded)
        elif file_type == "docx":
            return cls.parse_docx(decoded)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text for better analysis"""
        import re
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might break parsing
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
